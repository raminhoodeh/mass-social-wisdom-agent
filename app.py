import os
from typing import Optional, List, Dict
import re
import shutil
import threading
import uuid
import json
import requests
from datetime import datetime
from flask import Flask, render_template, jsonify, request, send_file, Response
from PIL import Image
from google import genai
from dotenv import load_dotenv
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────────────────────────
# Configuration
# ─────────────────────────────────────────────────────────────────
load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
SOCIAVAULT_API_KEY = os.getenv("SOCIAVAULT_API_KEY")

try:
    gemini_client = genai.Client()
except Exception as e:
    gemini_client = None
    print(f"Warning: Could not initialise Gemini Client: {e}")

BASE_DIR      = os.path.dirname(os.path.abspath(__file__))
SCAN_DIR      = os.path.join(BASE_DIR, "Scan")
RESULTS_DIR   = os.path.join(BASE_DIR, "Results")
PROCESSED_DIR = os.path.join(BASE_DIR, "Processed")

IMAGE_EXTENSIONS = ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp')

# ── FIXED: x-api-key (lowercase) is the working auth format ───────
SOCIAVAULT_BASE    = "https://api.sociavault.com/v1/scrape"
SOCIAVAULT_HEADERS = {"x-api-key": SOCIAVAULT_API_KEY}

CATEGORIES = [
    "Health & Wellbeing",
    "AI or Technology Advice",
    "Finance or Trading Advice or Tools",
    "Film or Movies or TV Shows",
    "Personal Branding or UGC or Social Media Tips",
    "Conspiracy Theories or Esoteric",
    "Romantic Relationships",
    "Other",
]

# Keywords that force a category regardless of Gemini's choice
FINANCE_KEYWORDS = {"polymarket", "kalshi", "mirofish", "trading", "stock", "forex", "crypto", "bitcoin", "options", "market"}
ROMANCE_KEYWORDS = {"dating", "relationship", "love", "girlfriend", "boyfriend", "marriage", "attraction", "romance", "rizz", "women", "men", "husband", "wife"}


# ─────────────────────────────────────────────────────────────────
# Global job tracker
# ─────────────────────────────────────────────────────────────────
jobs = {}  # job_id -> { status, log, result_path, failed_urls }

def ensure_folders():
    for d in [SCAN_DIR, RESULTS_DIR, PROCESSED_DIR]:
        os.makedirs(d, exist_ok=True)


# ─────────────────────────────────────────────────────────────────
# URL Detection
# ─────────────────────────────────────────────────────────────────
URL_RE = re.compile(
    r'https?://'
    r'(?:[-\w]+\.)+[a-zA-Z]{2,}'
    r'(?:/[^\s]*)?',
    re.IGNORECASE
)

def extract_urls(text: str) -> List[str]:
    """Pull all valid URLs out of a blob of messy text."""
    return [m.group(0).rstrip('.,;)"\'') for m in URL_RE.finditer(text)]

from urllib.parse import urlparse, parse_qs, urlencode, urlunparse

# Instagram paths that are NOT posts — skip them immediately
_IG_JUNK_PATHS = ("/accounts/", "/explore/", "/stories/", "/direct/", "/tv/")
_IG_POST_PATHS = ("/p/", "/reel/", "/reels/")

def sanitise_instagram_url(url: str) -> str:
    """
    Strip tracking/session parameters (igsh, si, app, feature, utm_*, etc.)
    from an Instagram URL so SociaVault receives a clean canonical URL.
    """
    keep_params = set()  # we keep nothing for Instagram — shortcode is in the path
    p = urlparse(url)
    return urlunparse(p._replace(query="", fragment=""))

def normalise_url(url: str) -> str:
    """
    Pre-process a raw URL before classification/API calls:
      1. Unwrap Instagram /accounts/login/?next=<actual_url> redirects
      2. Normalise /reels/ → /reel/ (SociaVault expects /reel/)
      3. Strip tracking params (igsh, si, app, feature, utm_*, etc.)
      4. For YouTube, keep only the 'v' param
    """
    from urllib.parse import unquote

    url = url.strip()

    # ── 1. Unwrap IG login redirects ──
    p = urlparse(url)
    if 'instagram.com' in p.netloc and p.path.startswith('/accounts/login'):
        qs = parse_qs(p.query)
        next_url = qs.get('next', [''])[0]
        if next_url:
            url = unquote(next_url)    # decode %2F etc.
            p = urlparse(url)

    # ── 2. Normalise /reels/ → /reel/ ──
    if 'instagram.com' in p.netloc:
        new_path = p.path.replace('/reels/', '/reel/')
        url = urlunparse(p._replace(path=new_path, query='', fragment=''))
        return url

    # ── 3. YouTube: keep only 'v' param, strip everything else ──
    if 'youtube.com' in p.netloc or 'youtu.be' in p.netloc:
        # Normalise m.youtube.com → www.youtube.com
        host = p.netloc.replace('m.youtube.com', 'www.youtube.com')
        qs = parse_qs(p.query)
        # For /shorts/ URLs the video ID is in the path, no params needed
        if '/shorts/' in p.path:
            return urlunparse(p._replace(netloc=host, query='', fragment=''))
        # For watch URLs keep only v=
        v = qs.get('v', [''])[0]
        new_query = f'v={v}' if v else ''
        return urlunparse(p._replace(netloc=host, query=new_query, fragment=''))

    return url


def classify_url(url: str) -> str:
    """Return 'instagram_reel', 'instagram_post', 'youtube', or 'unknown'.
    Rejects junk Instagram URLs (login redirects, profiles, homepage).
    URL should have been normalised first.
    """
    p = urlparse(url.lower())
    host = p.netloc
    path = p.path

    if 'instagram.com' in host:
        # Reject junk paths
        if not path or path == '/':
            return 'unknown'          # bare homepage
        if any(path.startswith(j) for j in _IG_JUNK_PATHS):
            return 'unknown'          # login redirect, stories, etc.
        if '/reel' in path:           # /reel/ and /reels/
            return 'instagram_reel'
        if '/p/' in path:             # standard post
            return 'instagram_post'
        return 'unknown'              # profile page, etc.

    if 'youtube.com' in host or 'youtu.be' in host:
        return 'youtube'

    return 'unknown'




# ─────────────────────────────────────────────────────────────────
# SociaVault helpers — with verbose error logging
# ─────────────────────────────────────────────────────────────────
def sv_get(endpoint: str, params: dict, timeout: int = 60):
    """Make a GET request to SociaVault and return the JSON body or raise."""
    url = f"{SOCIAVAULT_BASE}/{endpoint}"
    r = requests.get(url, headers=SOCIAVAULT_HEADERS, params=params, timeout=timeout)
    if not r.ok:
        raise Exception(f"HTTP {r.status_code}: {r.text[:300]}")
    return r.json()

def extract_transcript_text(data: dict) -> Optional[str]:
    """
    Parse transcript from SociaVault response.
    Instagram shape: { data: { data: { transcripts: { "0": { text } } } } }
    YouTube  shape:  { data: { transcript: { "0": { text } } } }
    Both use dict-keyed entries ("0", "1", …) with a "text" field.
    """
    try:
        inner = data.get("data", {})
        if not isinstance(inner, dict):
            return None

        # Check every plausible location (singular + plural, different depths)
        candidates = [
            inner.get("transcript"),                             # data.data.transcript  (YouTube)
            inner.get("transcripts"),                            # data.data.transcripts (IG variant)
            inner.get("data", {}).get("transcript") if isinstance(inner.get("data"), dict) else None,
            inner.get("data", {}).get("transcripts") if isinstance(inner.get("data"), dict) else None,
        ]

        for transcripts in candidates:
            if not transcripts or not isinstance(transcripts, dict):
                continue
            all_texts = []
            for key in sorted(transcripts.keys(), key=lambda k: int(k) if k.isdigit() else 0):
                t = transcripts[key]
                if isinstance(t, dict):
                    txt = t.get("text") or t.get("transcript") or ""
                    if txt and isinstance(txt, str):
                        all_texts.append(txt.strip())
            if all_texts:
                return " ".join(all_texts)   # join with space for natural flow

        return None
    except Exception:
        return None


def _walk_edges(edges_raw) -> Optional[str]:
    """Extract text from edges regardless of whether it's a dict or list."""
    if isinstance(edges_raw, dict) and edges_raw:
        texts = []
        for k in sorted(edges_raw.keys(), key=lambda x: int(x) if x.isdigit() else 0):
            t = edges_raw[k].get("node", {}).get("text", "") if isinstance(edges_raw[k], dict) else ""
            if t: texts.append(t.strip())
        return "\n".join(texts) if texts else None
    elif isinstance(edges_raw, list) and edges_raw:
        texts = [e.get("node", {}).get("text", "") for e in edges_raw if isinstance(e, dict)]
        texts = [t.strip() for t in texts if t]
        return "\n".join(texts) if texts else None
    return None

def extract_caption_text(data: dict) -> Optional[str]:
    """
    Parse caption from SociaVault post-info response.
    Actual path: data["data"]["data"]["xdt_shortcode_media"]["edge_media_to_caption"]["edges"]
    where edges is a dict with string-number keys {"0": {node: {text: ...}}}
    Tries multiple nesting depths as SociaVault can vary.
    """
    try:
        # Try all possible nesting depths for robustness
        candidates = [
            data.get("data", {}).get("data", {}).get("xdt_shortcode_media", {}),        # 2-level (confirmed)
            data.get("data", {}).get("data", {}).get("data", {}).get("xdt_shortcode_media", {}),  # 3-level (backup)
        ]
        for media in candidates:
            if not media:
                continue
            edges_raw = media.get("edge_media_to_caption", {}).get("edges", {})
            result = _walk_edges(edges_raw)
            if result:
                return result
            # Direct fallback fields
            for field in ("caption", "title"):
                v = media.get(field)
                if v and isinstance(v, str) and v.strip():
                    return v.strip()
        return None
    except Exception:
        return None


def fetch_instagram_transcript(post_url: str, log_fn) -> Optional[str]:
    try:
        clean_url = sanitise_instagram_url(post_url)
        log_fn("    → Calling SociaVault transcript API…")
        data = sv_get("instagram/transcript", {"url": clean_url}, timeout=90)
        text = extract_transcript_text(data)
        if text:
            log_fn(f"    ✅ Transcript received ({len(text)} chars)")
        else:
            log_fn(f"    ⚠️ Transcript API returned empty — raw keys: {list(data.keys())}")
        return text
    except Exception as e:
        log_fn(f"    ❌ Transcript API error: {e}")
        return None

def fetch_instagram_post_info(post_url: str, log_fn) -> dict:
    """
    Returns enriched dict:
      caption    – cleaned caption string or None
      is_video   – True/False
      typename   – XDTGraphVideo / XDTGraphImage / XDTGraphSidecar
      image_urls – list of CDN image URLs (display_url for single; per-slide for carousel)
      _raw       – full API response
    """
    clean_url = sanitise_instagram_url(post_url)
    try:
        log_fn("    → Calling SociaVault post-info API…")
        data = sv_get("instagram/post-info", {"url": clean_url}, timeout=60)


        # Resolve media object (try 2-level and 3-level nesting)
        media = (
            data.get("data", {}).get("data", {}).get("xdt_shortcode_media")
            or data.get("data", {}).get("data", {}).get("data", {}).get("xdt_shortcode_media")
            or {}
        )

        caption   = extract_caption_text(data)
        is_video  = bool(media.get("is_video"))
        typename  = media.get("__typename", "")

        # Collect slide image URLs
        image_urls: List[str] = []
        sidecar = media.get("edge_sidecar_to_children", {})
        if sidecar:
            edges = sidecar.get("edges", {})
            nodes = list(edges.values()) if isinstance(edges, dict) else edges
            for n in nodes:
                node = n.get("node", {}) if isinstance(n, dict) else {}
                if not node.get("is_video"):
                    url_img = node.get("display_url") or ""
                    if url_img:
                        image_urls.append(url_img)
        elif not is_video:
            # Single static image
            url_img = media.get("display_url") or ""
            if url_img:
                image_urls.append(url_img)

        if caption:
            log_fn(f"    ✅ Caption received ({len(caption)} chars)")
        else:
            log_fn(f"    ⚠️ No caption found for this post")

        log_fn(f"    ℹ️  Type: {typename or ('video' if is_video else 'image')}, "
               f"image slides: {len(image_urls)}")

        return {
            "caption":    caption,
            "is_video":   is_video,
            "typename":   typename,
            "image_urls": image_urls,
            "_raw":       data,
        }
    except Exception as e:
        log_fn(f"    ❌ Post-info API error: {e}")
        return {"caption": None, "is_video": False, "typename": "", "image_urls": []}


def fetch_youtube_transcript(video_url: str, log_fn) -> Optional[str]:
    try:
        log_fn("    → Calling SociaVault YouTube transcript API…")
        data = sv_get("youtube/video/transcript", {"url": video_url}, timeout=90)

        # Try the structured parser first (handles dict-keyed transcript objects)
        text = extract_transcript_text(data)

        if not text:
            # Fallback chain — ONLY accept plain strings, never dicts
            candidates = [
                data.get("transcript"),
                data.get("text"),
                data.get("data", {}).get("transcript"),
                data.get("data", {}).get("text"),
                data.get("data", {}).get("data", {}).get("transcript"),
                data.get("data", {}).get("data", {}).get("text"),
            ]
            for c in candidates:
                if c and isinstance(c, str) and c.strip():
                    text = c.strip()
                    break

        if text:
            log_fn(f"    ✅ YouTube transcript received ({len(text)} chars)")
        else:
            log_fn(f"    ⚠️ YouTube transcript empty — raw keys: {list(data.keys())}")
        return text
    except Exception as e:
        log_fn(f"    ❌ YouTube transcript API error: {e}")
        return None



# ─────────────────────────────────────────────────────────────────
# Gemini helpers
# ─────────────────────────────────────────────────────────────────
def gemini_text(prompt: str) -> str:
    if gemini_client is None:
        return "[Gemini unavailable]"
    response = gemini_client.models.generate_content(
        model="gemini-2.5-flash",
        contents=[prompt]
    )
    return response.text.strip()

def gemini_vision(image: Image.Image, prompt: str) -> str:
    if gemini_client is None:
        return "[Gemini unavailable]"
    response = gemini_client.models.generate_content(
        model="gemini-2.5-flash",
        contents=[image, prompt]
    )
    return response.text.strip()

def compose_output(transcript: Optional[str], caption: Optional[str],
                   ocr_slides: Optional[List[str]] = None,
                   kind: str = "", lenient: bool = False) -> str:
    """
    Combine transcript, caption, and/or OCR slide text into a single polished,
    readable paragraph using Gemini. No labels like 'Caption:' or 'Slide 1:'.

    lenient=True uses a less strict prompt — invoked by the agent retry loop
    when the standard pass scores below quality threshold.
    """
    has_transcript = bool(transcript and transcript.strip())
    has_caption    = bool(caption and caption.strip())
    has_slides     = bool(ocr_slides)

    # Build the raw material block for Gemini
    raw_parts = []
    if has_transcript:
        raw_parts.append(f"TRANSCRIPT:\n{transcript.strip()}")
    if has_slides:
        for i, s in enumerate(ocr_slides, 1):
            raw_parts.append(f"SLIDE {i} TEXT (OCR):\n{s}")
    if has_caption:
        raw_parts.append(f"CAPTION:\n{caption.strip()}")

    if not raw_parts:
        return ""

    raw_block = "\n\n".join(raw_parts)

    if kind == "youtube":
        context_note = "This content is from a YouTube video."
        cap_note = ("The caption/description is supplementary — summarise it briefly if it adds value, "
                    "but skip it if it\'s a generic channel description.")
    elif kind in ("instagram_reel", "instagram_post", "instagram_video"):
        context_note = "This content is from an Instagram post."
        cap_note = "The caption is important context — weave it naturally into the output."
    else:
        context_note = ""
        cap_note = ""

    if lenient:
        rules = (
            "Rules (lenient mode — maximum content preservation):\n"
            "- Your primary goal is faithful reproduction. Keep ALL details, even if phrasing is rough.\n"
            "- Apply only minimal grammar fixes — do NOT restructure, shorten, or summarise.\n"
            "- If only a caption is available, present it clearly as the main content.\n"
            "- Do NOT discard content for being 'redundant' or 'promotional' — include everything.\n"
            "- Output ONLY the final text. No preamble, no metadata."
        )
    else:
        rules = (
            f"Rules:\n"
            f"- Do NOT aggressively summarise the content. Keep the final output as detailed and close to the original length as possible.\n"
            f"- Preserve all original details, anecdotes, examples, and nuances.\n"
            f"- Fix grammar, capitalisation, and punctuation throughout to improve flow, but do NOT strip out interesting information.\n"
            f"- Do NOT use labels like 'Caption:', 'Transcript:', 'Slide 1:', etc. in the output.\n"
            f"- Write in natural, flowing paragraph prose.\n"
            f"- If there are multiple slides, merge their content into a natural narrative flow — do not number them.\n"
            f"{cap_note}\n"
            f"- Output ONLY the final polished text. No preamble, no metadata."
        )

    prompt = f"""You are a content editor preparing notes for a personal knowledge base.
{context_note}

Your task: combine the following raw extracted content into a single, cohesive, readable passage.

{rules}

RAW CONTENT:
{raw_block}"""

    try:
        return gemini_text(prompt)
    except Exception:
        # Fallback: join raw parts without labels
        fallback = []
        if transcript: fallback.append(transcript.strip())
        if ocr_slides: fallback.extend(ocr_slides)
        if caption:    fallback.append(caption.strip())
        return "\n\n".join(fallback)

def categorise_content(content: str, source_url: str) -> str:
    """
    Pick the best category for this content.
    Keyword overrides are checked first for Finance and Romantic Relationships,
    then Gemini is called for everything else.
    """
    content_lower = content.lower()

    # Hard keyword overrides (checked before Gemini)
    # Using regex word boundaries so "comment" doesn't match "men", etc.
    for kw in FINANCE_KEYWORDS:
        if re.search(r'\b' + re.escape(kw) + r'\b', content_lower):
            return "Finance or Trading Advice or Tools"

    for kw in ROMANCE_KEYWORDS:
        if re.search(r'\b' + re.escape(kw) + r'\b', content_lower):
            return "Romantic Relationships"

    cats_list = "\n".join(f"{i+1}. {c}" for i, c in enumerate(CATEGORIES))
    prompt = (
        f"You are a content categorisation engine. Given the following text extracted from "
        f"'{source_url}', assign it to EXACTLY ONE of these categories by returning only the "
        f"category name, nothing else:\n\n{cats_list}\n\nContent:\n{content}"
    )
    result = gemini_text(prompt)
    result_lower = result.lower()
    for cat in CATEGORIES:
        if cat.lower() in result_lower or result_lower in cat.lower():
            return cat
    return "Other"


def assess_output_quality(content: str) -> int:
    """
    Agent self-assessment step: ask Gemini to score the quality of composed
    output on a scale of 1–5. Used in the agent retry loop to decide whether
    to re-compose with a more lenient strategy.

    Returns an integer 1–5:
      1 = Gibberish, error message, or empty/useless
      2 = Very sparse — just a short caption or a few disconnected words
      3 = Acceptable but thin
      4 = Good — coherent and informative
      5 = Excellent — rich, detailed, well-structured
    """
    if not content or len(content.strip()) < 80:
        return 1  # trivially too short — skip the API call

    prompt = (
        "Rate the quality of the following extracted text on a scale from 1 to 5. "
        "Return ONLY the single digit (1, 2, 3, 4, or 5) — nothing else.\n\n"
        "1 = Gibberish, error message, or completely empty/useless.\n"
        "2 = Very sparse — just a short caption or a few disconnected words.\n"
        "3 = Acceptable but thin — some real content, but limited detail.\n"
        "4 = Good — coherent, informative prose with meaningful detail.\n"
        "5 = Excellent — rich, detailed, well-structured content.\n\n"
        f"Text to evaluate (first 800 chars):\n{content[:800]}"
    )
    try:
        result = gemini_text(prompt).strip()
        for ch in result:
            if ch.isdigit():
                return max(1, min(5, int(ch)))
        return 3  # neutral default if parsing fails
    except Exception:
        return 3


def sort_items_by_similarity(items: List[Dict], log_fn) -> List[Dict]:
    """
    Within each category, use Gemini to reorder items so that
    topically similar items are adjacent. Returns a new list in
    category-grouped, similarity-sorted order.
    """
    if len(items) <= 1:
        return items

    # Group by category
    grouped: Dict[str, List[Dict]] = {cat: [] for cat in CATEGORIES}
    for item in items:
        cat = item.get("category", "Other")
        if cat not in grouped:
            cat = "Other"
        grouped[cat].append(item)

    sorted_items: List[Dict] = []

    for cat in CATEGORIES:
        entries = grouped[cat]
        if len(entries) <= 1:
            sorted_items.extend(entries)
            continue

        # Build a concise summary of each entry for Gemini
        summaries = []
        for i, entry in enumerate(entries):
            # Take first 200 chars of content as a topic fingerprint
            snippet = entry["content"][:200].replace("\n", " ").strip()
            summaries.append(f"{i}: {snippet}")

        summaries_block = "\n".join(summaries)

        prompt = (
            f"You have {len(entries)} content items in the '{cat}' category. "
            f"Reorder them so that items about similar topics/themes are adjacent to each other "
            f"(i.e. cluster by topic similarity).\n\n"
            f"Items:\n{summaries_block}\n\n"
            f"Return ONLY a comma-separated list of the item numbers in the new order. "
            f"Example for 5 items: 3,1,0,4,2\n"
            f"Do not add any other text."
        )

        try:
            result = gemini_text(prompt).strip()
            # Parse the indices
            new_order = []
            for part in result.replace(" ", "").split(","):
                part = part.strip()
                if part.isdigit():
                    idx = int(part)
                    if 0 <= idx < len(entries) and idx not in new_order:
                        new_order.append(idx)
            # Add any missing indices at the end (safety)
            for i in range(len(entries)):
                if i not in new_order:
                    new_order.append(i)

            sorted_entries = [entries[i] for i in new_order]
            log_fn(f"  🔀 '{cat}': reordered {len(entries)} items by topic similarity")
            sorted_items.extend(sorted_entries)
        except Exception as e:
            log_fn(f"  ⚠️ Similarity sort failed for '{cat}': {e} — keeping original order")
            sorted_items.extend(entries)

    return sorted_items


def ocr_image_local(filepath: str) -> str:
    """Run Gemini vision OCR on a local image file."""
    with Image.open(filepath) as img:
        prompt = (
            "Extract ALL visible text from this image exactly as it appears. "
            "Preserve line breaks and structure. Output plain text only."
        )
        return gemini_vision(img, prompt)

def ocr_image_url(image_url: str, log_fn, label: str = "") -> Optional[str]:
    """
    Download an image from a CDN URL into memory and OCR it with Gemini Vision.
    Returns extracted text or None.
    """
    try:
        log_fn(f"    🖼️  OCR-ing{' ' + label if label else ''} image…")
        resp = requests.get(image_url, timeout=30)
        resp.raise_for_status()
        img = Image.open(__import__('io').BytesIO(resp.content))
        prompt = (
            "Extract ALL visible text from this image exactly as it appears. "
            "Preserve line breaks and structure. Output plain text only. "
            "If there is no meaningful text, return an empty string."
        )
        text = gemini_vision(img, prompt)
        if text and text.strip():
            log_fn(f"    ✅ OCR extracted {len(text)} chars")
            return text.strip()
        else:
            log_fn(f"    ⚠️ No text found in image")
            return None
    except Exception as e:
        log_fn(f"    ❌ OCR failed: {e}")
        return None


# ─────────────────────────────────────────────────────────────────
# Extraction pipeline
# ─────────────────────────────────────────────────────────────────
def extract_from_url(url: str, log_fn) -> Optional[Dict]:
    """
    Given a URL, return a dict: { url, category, content }
    or None if extraction completely failed.

    Routing logic:
      instagram /reel/  → transcript + caption  → compose_output
      instagram /p/     → post-info sub-type:
                            XDTGraphSidecar (carousel) → OCR slides + caption → compose_output
                            XDTGraphVideo              → transcript + caption → compose_output
                            XDTGraphImage (single img) → OCR + caption        → compose_output
      youtube           → transcript → compose_output
    """
    # Normalise the URL first (unwrap redirects, strip tracking params, etc.)
    url = normalise_url(url)
    kind = classify_url(url)
    log_fn(f"🔗 Processing ({kind}): {url}")


    transcript: Optional[str] = None
    caption:    Optional[str] = None
    ocr_slides: List[str]     = []

    if kind == "instagram_reel":
        post_info  = fetch_instagram_post_info(url, log_fn)
        caption    = post_info.get("caption") or None
        transcript = fetch_instagram_transcript(url, log_fn)

    elif kind == "instagram_post":
        post_info  = fetch_instagram_post_info(url, log_fn)
        caption    = post_info.get("caption") or None
        typename   = post_info.get("typename", "")
        is_video   = post_info.get("is_video", False)
        image_urls = post_info.get("image_urls", [])

        if typename == "XDTGraphSidecar" or (not is_video and len(image_urls) > 1):
            log_fn(f"    🏗️  Carousel detected — OCR-ing {len(image_urls)} slide(s)…")
            for i, img_url in enumerate(image_urls):
                text = ocr_image_url(img_url, log_fn, label=f"slide {i+1}/{len(image_urls)}")
                if text:
                    ocr_slides.append(text)
                # Pace API requests to avoid Gemini Free Tier RPM (15 Requests Per Minute) block
                if i < len(image_urls) - 1:
                    time.sleep(3)

        elif is_video or typename == "XDTGraphVideo":
            log_fn("    🎬 Video post — fetching transcript…")
            transcript = fetch_instagram_transcript(url, log_fn)

        else:
            if image_urls:
                log_fn("    🖼️  Single image post — running OCR…")
                text = ocr_image_url(image_urls[0], log_fn, label="post image")
                if text:
                    ocr_slides.append(text)
            else:
                log_fn("    ⚠️ Could not determine image URL for this post")

    elif kind == "youtube":
        transcript = fetch_youtube_transcript(url, log_fn)

    else:
        log_fn(f"  ⚠️ Unrecognised URL type, skipping.")
        return None

    # Check we have something to work with
    if not transcript and not caption and not ocr_slides:
        log_fn("  ⚠️ No content could be extracted from this URL.")
        return None

    # Compose into polished prose
    log_fn("  ✍️  Composing polished output…")
    final_text = compose_output(transcript, caption, ocr_slides or None, kind=kind)

    if not final_text or not final_text.strip():
        log_fn("  ⚠️ Output composition returned empty.")
        return None

    # ── Agent self-assessment & retry loop ──────────────────────
    log_fn("  🤖 Agent: assessing output quality…")
    quality_score = assess_output_quality(final_text)
    log_fn(f"  🤖 Agent: quality score {quality_score}/5")

    if quality_score < 3:
        log_fn(f"  🔄 Agent decision: score too low ({quality_score}/5) — retrying with lenient composition…")
        retry_text  = compose_output(transcript, caption, ocr_slides or None, kind=kind, lenient=True)
        retry_score = assess_output_quality(retry_text) if retry_text and retry_text.strip() else 0
        log_fn(f"  🤖 Agent: retry quality score {retry_score}/5")
        if retry_score > quality_score and retry_text.strip():
            log_fn(f"  ✅ Agent: retry improved output ({quality_score} → {retry_score}) — using retry result")
            final_text = retry_text
        else:
            log_fn(f"  ⚠️ Agent: retry did not improve quality — keeping original output")
    else:
        log_fn(f"  ✅ Agent: quality acceptable — proceeding")
    # ────────────────────────────────────────────────────────────

    try:
        category = categorise_content(final_text, url)
        log_fn(f"  ✅ Categorised as: {category}")
    except Exception as e:
        log_fn(f"  ⚠️ Categorisation failed ({e}), defaulting to Other.")
        category = "Other"

    return {"url": url, "category": category, "content": final_text, "caption": caption}


def extract_from_local_image(filepath: str, filename: str, log_fn) -> Optional[Dict]:
    """OCR a local image file and return the same dict shape."""
    log_fn(f"🖼️  Processing local image: {filename}")
    try:
        text = ocr_image_local(filepath)
        if not text or not text.strip():
            log_fn(f"  ⚠️ No text found in {filename}")
            return None
        category = categorise_content(text, filename)
        log_fn(f"  ✅ Categorised as: {category}")

        dest = os.path.join(PROCESSED_DIR, filename)
        counter = 1
        while os.path.exists(dest):
            base, ext = os.path.splitext(filename)
            dest = os.path.join(PROCESSED_DIR, f"{base}_{counter}{ext}")
            counter += 1
        shutil.move(filepath, dest)

        return {"url": filename, "category": category, "content": text, "caption": None}
    except Exception as e:
        log_fn(f"  ❌ Error on {filename}: {e}")
        return None


# ─────────────────────────────────────────────────────────────────
# Word document generation
# ─────────────────────────────────────────────────────────────────
def generate_docx(items: List[Dict], output_path: str):
    """
    Build a clean, Notion-importable Word document grouped by category.
    Each entry: embedded source link + plain paragraph text.
    """
    doc = Document()

    title = doc.add_heading("Mass Social Wisdom Agent — Session Results", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    date_para = doc.add_paragraph(
        datetime.now().strftime("Extracted on %d %B %Y at %H:%M")
    )
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x99)
    doc.add_paragraph()

    grouped: Dict[str, List[Dict]] = {cat: [] for cat in CATEGORIES}
    for item in items:
        cat = item.get("category", "Other")
        if cat not in grouped:
            cat = "Other"
        grouped[cat].append(item)

    for cat in CATEGORIES:
        entries = grouped[cat]
        if not entries:
            continue

        doc.add_heading(cat, level=1)

        for entry in entries:
            source_url = entry["url"]
            content    = entry["content"]

            link_para = doc.add_paragraph()
            link_run  = link_para.add_run(f"Source: {source_url}")
            link_run.font.color.rgb = RGBColor(0x22, 0x66, 0xCC)
            link_run.font.size = Pt(9)

            if entry.get("caption"):
                cap_para = doc.add_paragraph()
                cap_para.add_run("Caption: ").bold = True
                cap_para.add_run(entry["caption"].strip())
                cap_para.paragraph_format.space_after = Pt(12)

            for block in content.split("\n\n"):
                block = block.strip()
                if block:
                    p = doc.add_paragraph(block)
                    p.paragraph_format.space_after = Pt(6)

            doc.add_paragraph("─" * 60).runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.save(output_path)


# ─────────────────────────────────────────────────────────────────
# Background job runner — with retry support
# ─────────────────────────────────────────────────────────────────
def run_extraction_job(job_id: str, urls: List[str], include_scan_folder: bool):
    job = jobs[job_id]
    log = job["log"]

    def log_fn(msg: str):
        log.append(msg)
        print(msg)

    def is_stopped() -> bool:
        return job.get("stop_requested", False)

    log_fn("🚀 Mass Social Wisdom Agent started.")
    items:       List[Dict] = []
    failed_urls: List[str]  = []

    # ── 1. Local images ──────────────────────────────────────────
    if include_scan_folder:
        ensure_folders()
        try:
            local_files = [
                f for f in os.listdir(SCAN_DIR)
                if f.lower().endswith(IMAGE_EXTENSIONS)
            ]
        except Exception:
            local_files = []

        if local_files:
            log_fn(f"\n📂 Found {len(local_files)} image(s) in Scan folder.")
            for filename in local_files:
                result = extract_from_local_image(
                    os.path.join(SCAN_DIR, filename), filename, log_fn
                )
                if result:
                    items.append(result)
                    job["live_items"].append(result)
                else:
                    failed_urls.append(filename)
        else:
            log_fn("📂 Scan folder is empty — skipping local images.")

    # ── 2. URLs (with per-URL error isolation) ─────────────────
    if urls:
        log_fn(f"\n🌐 Processing {len(urls)} URL(s)…")
        for i, url in enumerate(urls):
            if is_stopped():
                log_fn(f"\n🛑 Stop requested — {len(urls) - i} URL(s) remaining.")
                # Add remaining as failed so user can retry
                for remaining_url in urls[i:]:
                    failed_urls.append(remaining_url)
                break
            try:
                # Pace the main loop to avoid hitting Gemini 15 RPM Free Tier limits
                if i > 0:
                    time.sleep(4)
                
                result = extract_from_url(url, log_fn)
                if result:
                    items.append(result)
                    job["live_items"].append(result)
                else:
                    failed_urls.append(url)
            except Exception as e:
                log_fn(f"  ❌ Unhandled error on {url}: {e}")
                failed_urls.append(url)

    # ── 3. Report failures ────────────────────────────────────
    if failed_urls:
        log_fn(f"\n⚠️ {len(failed_urls)} URL(s) produced no content:")
        for u in failed_urls:
            log_fn(f"  • {u}")
        # Save failed URLs to a text file for easy retry
        failed_path = os.path.join(RESULTS_DIR, f"failed_urls_{job_id[:8]}.txt")
        with open(failed_path, "w", encoding="utf-8") as f:
            f.write("\n".join(failed_urls))
        log_fn(f"  → Saved to Results/failed_urls_{job_id[:8]}.txt — paste these back to retry.")
        job["failed_urls"]  = failed_urls
        job["failed_path"]  = failed_path
        job["failed_name"]  = f"failed_urls_{job_id[:8]}.txt"

    # ── 4. Similarity-sort within each category ────────────────
    if not items:
        log_fn("\n⚠️ Nothing was successfully extracted — no document generated.")
        job["status"] = "done"
        return

    log_fn("\n🔀 Sorting items by topic similarity within each category…")
    items = sort_items_by_similarity(items, log_fn)
    # Update live_items so the frontend re-renders in sorted order
    job["live_items"] = items
    job["items_sorted"] = True

    # ── 5. Generate document ──────────────────────────────────
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    doc_name  = f"Social_Extractor_{timestamp}.docx"
    doc_path  = os.path.join(RESULTS_DIR, doc_name)

    log_fn(f"\n📄 Generating Word document ({len(items)} item(s))…")
    try:
        generate_docx(items, doc_path)

        log_fn(f"✅ Done! Saved: Results/{doc_name}")
        job["status"]      = "done"
        job["result_path"] = doc_path
        job["result_name"] = doc_name
    except Exception as e:
        log_fn(f"❌ Failed to generate document: {e}")
        job["status"] = "error"


# ─────────────────────────────────────────────────────────────────
# Flask app
# ─────────────────────────────────────────────────────────────────
app = Flask(__name__)

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/favicon.ico")
def favicon():
    return Response("", status=204)

@app.route("/extract", methods=["POST"])
def start_extraction():
    data        = request.json or {}
    raw_text    = data.get("urls_text", "")
    include_scan = data.get("include_scan", True)

    urls = extract_urls(raw_text)
    urls = list(dict.fromkeys(urls))  # dedupe preserving order

    if not urls and not include_scan:
        return jsonify({"success": False, "message": "No valid URLs found in the input."})

    ensure_folders()
    job_id = str(uuid.uuid4())
    jobs[job_id] = {
        "status":        "running",
        "log":           [],
        "live_items":    [],
        "result_path":   None,
        "result_name":   None,
        "failed_urls":   [],
        "failed_name":   None,
        "stop_requested": False,
    }


    thread = threading.Thread(
        target=run_extraction_job,
        args=(job_id, urls, include_scan),
        daemon=True
    )
    thread.start()

    return jsonify({"success": True, "job_id": job_id, "url_count": len(urls)})

@app.route("/status/<job_id>")
def job_status(job_id):
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Unknown job"}), 404
    return jsonify({
        "status":       job["status"],
        "log":          job["log"],
        "live_items":   job.get("live_items", []),
        "result_name":  job.get("result_name"),
        "failed_name":  job.get("failed_name"),
        "failed_count": len(job.get("failed_urls", [])),
    })


@app.route("/download/<job_id>")
def download_result(job_id):
    job = jobs.get(job_id)
    if not job or not job.get("result_path"):
        return jsonify({"error": "No result available"}), 404
    return send_file(job["result_path"], as_attachment=True, download_name=job["result_name"])

@app.route("/download-failed/<job_id>")
def download_failed(job_id):
    job = jobs.get(job_id)
    if not job or not job.get("failed_path"):
        return jsonify({"error": "No failed URLs file"}), 404
    return send_file(job["failed_path"], as_attachment=True, download_name=job["failed_name"])

@app.route("/stop/<job_id>", methods=["POST"])
def stop_job(job_id):
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Unknown job"}), 404
    job["stop_requested"] = True
    return jsonify({"success": True})

if __name__ == "__main__":
    ensure_folders()
    # Cloud Run provides the port via the PORT environment variable
    port = int(os.environ.get("PORT", 5001))
    app.run(host="0.0.0.0", port=port, debug=False)
